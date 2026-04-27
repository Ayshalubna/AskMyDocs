"""
Document processing service.
Handles parsing of PDF, DOCX, TXT, MD, CSV, HTML files.
Implements recursive character text splitter with smart chunking.
"""

import asyncio
import csv
import hashlib
import io
import re
from pathlib import Path
from typing import List, Optional, Tuple

import structlog
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

from app.core.config import settings
from app.core.exceptions import DocumentProcessingError, UnsupportedFileTypeError

logger = structlog.get_logger(__name__)


class DocumentProcessor:
    """Multi-format document processor with intelligent chunking."""

    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            separators=settings.CHUNK_SEPARATORS,
            length_function=len,
            is_separator_regex=False,
        )
        self._parsers = {
            ".pdf": self._parse_pdf,
            ".txt": self._parse_text,
            ".md": self._parse_markdown,
            ".docx": self._parse_docx,
            ".csv": self._parse_csv,
            ".html": self._parse_html,
        }

    async def process_file(
        self,
        file_path: Path,
        doc_id: str,
        filename: str,
        extra_metadata: Optional[dict] = None,
    ) -> Tuple[List[Document], dict]:
        """
        Process a file and return LangChain Documents with metadata.
        Returns (chunks, file_metadata)
        """
        extension = file_path.suffix.lower()

        if extension not in self._parsers:
            raise UnsupportedFileTypeError(extension, list(self._parsers.keys()))

        logger.info("processing_document", doc_id=doc_id, filename=filename, extension=extension)

        try:
            # Parse raw text + page info
            raw_documents, file_metadata = await asyncio.get_event_loop().run_in_executor(
                None, self._parsers[extension], file_path, doc_id, filename
            )

            # Split into chunks
            chunks = self.text_splitter.split_documents(raw_documents)

            # Enrich chunk metadata
            chunks = self._enrich_metadata(chunks, doc_id, filename, extra_metadata or {})

            # Deduplicate
            chunks = self._deduplicate_chunks(chunks)

            logger.info(
                "document_processed",
                doc_id=doc_id,
                num_chunks=len(chunks),
                num_pages=file_metadata.get("num_pages"),
            )
            return chunks, file_metadata

        except (UnsupportedFileTypeError, DocumentProcessingError):
            raise
        except Exception as e:
            logger.error("document_processing_failed", doc_id=doc_id, error=str(e))
            raise DocumentProcessingError(filename, str(e))

    def _parse_pdf(self, file_path: Path, doc_id: str, filename: str):
        """Parse PDF with page-level metadata."""
        try:
            import pypdf
        except ImportError:
            raise DocumentProcessingError(filename, "pypdf not installed")

        documents = []
        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            num_pages = len(reader.pages)

            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text() or ""
                text = self._clean_text(text)
                if text.strip():
                    documents.append(
                        Document(
                            page_content=text,
                            metadata={
                                "doc_id": doc_id,
                                "filename": filename,
                                "page_number": page_num,
                                "source": str(file_path),
                            },
                        )
                    )

        return documents, {"num_pages": num_pages}

    def _parse_docx(self, file_path: Path, doc_id: str, filename: str):
        """Parse DOCX preserving heading structure."""
        try:
            import docx
        except ImportError:
            raise DocumentProcessingError(filename, "python-docx not installed")

        doc = docx.Document(file_path)
        sections = []
        current_section = ""
        current_heading = "Introduction"

        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                if current_section.strip():
                    sections.append((current_heading, current_section))
                current_heading = para.text
                current_section = ""
            else:
                current_section += para.text + "\n"

        if current_section.strip():
            sections.append((current_heading, current_section))

        documents = []
        for heading, content in sections:
            cleaned = self._clean_text(content)
            if cleaned.strip():
                documents.append(
                    Document(
                        page_content=cleaned,
                        metadata={
                            "doc_id": doc_id,
                            "filename": filename,
                            "section": heading,
                            "source": str(file_path),
                        },
                    )
                )

        return documents, {"num_pages": None, "num_sections": len(sections)}

    def _parse_text(self, file_path: Path, doc_id: str, filename: str):
        """Parse plain text."""
        text = file_path.read_text(encoding="utf-8", errors="replace")
        text = self._clean_text(text)
        documents = [
            Document(
                page_content=text,
                metadata={"doc_id": doc_id, "filename": filename, "source": str(file_path)},
            )
        ]
        return documents, {"num_pages": None}

    def _parse_markdown(self, file_path: Path, doc_id: str, filename: str):
        """Parse Markdown, splitting by headers."""
        text = file_path.read_text(encoding="utf-8", errors="replace")
        
        # Split by headers
        sections = re.split(r"(^#{1,3} .+$)", text, flags=re.MULTILINE)
        
        documents = []
        current_heading = "Overview"
        for chunk in sections:
            if re.match(r"^#{1,3} ", chunk):
                current_heading = chunk.strip("# \n")
            elif chunk.strip():
                cleaned = self._clean_text(chunk)
                documents.append(
                    Document(
                        page_content=cleaned,
                        metadata={
                            "doc_id": doc_id,
                            "filename": filename,
                            "section": current_heading,
                            "source": str(file_path),
                        },
                    )
                )

        if not documents:
            documents, _ = self._parse_text(file_path, doc_id, filename)

        return documents, {"num_pages": None}

    def _parse_csv(self, file_path: Path, doc_id: str, filename: str):
        """Parse CSV, converting rows to natural language."""
        documents = []
        with open(file_path, newline="", encoding="utf-8", errors="replace") as f:
            reader = csv.DictReader(f)
            rows = list(reader)
            headers = reader.fieldnames or []

            # Group rows into chunks of 20 for context
            chunk_size = 20
            for i in range(0, len(rows), chunk_size):
                batch = rows[i : i + chunk_size]
                text_parts = [f"CSV Data ({filename}) - Rows {i+1} to {i+len(batch)}:"]
                text_parts.append(f"Columns: {', '.join(headers)}\n")
                
                for row_num, row in enumerate(batch, i + 1):
                    row_text = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
                    text_parts.append(f"Row {row_num}: {row_text}")
                
                documents.append(
                    Document(
                        page_content="\n".join(text_parts),
                        metadata={
                            "doc_id": doc_id,
                            "filename": filename,
                            "row_start": i + 1,
                            "row_end": i + len(batch),
                            "source": str(file_path),
                        },
                    )
                )

        return documents, {"num_rows": len(rows), "num_columns": len(headers)}

    def _parse_html(self, file_path: Path, doc_id: str, filename: str):
        """Parse HTML, stripping tags."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise DocumentProcessingError(filename, "beautifulsoup4 not installed")

        html = file_path.read_text(encoding="utf-8", errors="replace")
        soup = BeautifulSoup(html, "html.parser")
        
        # Remove scripts and styles
        for tag in soup(["script", "style", "nav", "footer"]):
            tag.decompose()
        
        text = soup.get_text(separator="\n")
        text = self._clean_text(text)
        
        documents = [
            Document(
                page_content=text,
                metadata={
                    "doc_id": doc_id,
                    "filename": filename,
                    "title": soup.title.string if soup.title else filename,
                    "source": str(file_path),
                },
            )
        ]
        return documents, {"num_pages": None}

    def _clean_text(self, text: str) -> str:
        """Normalize whitespace and remove junk characters."""
        # Replace multiple newlines with double newline
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Replace multiple spaces
        text = re.sub(r" {2,}", " ", text)
        # Remove null bytes and control characters
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
        return text.strip()

    def _enrich_metadata(
        self,
        chunks: List[Document],
        doc_id: str,
        filename: str,
        extra_metadata: dict,
    ) -> List[Document]:
        """Add chunk index and content hash to each chunk."""
        for i, chunk in enumerate(chunks):
            content_hash = hashlib.md5(chunk.page_content.encode()).hexdigest()[:8]
            chunk.metadata.update(
                {
                    "doc_id": doc_id,
                    "filename": filename,
                    "chunk_index": i,
                    "chunk_id": f"{doc_id}_{i}_{content_hash}",
                    "char_count": len(chunk.page_content),
                    **extra_metadata,
                }
            )
        return chunks

    def _deduplicate_chunks(self, chunks: List[Document]) -> List[Document]:
        """Remove near-duplicate chunks based on content hash."""
        seen_hashes = set()
        unique_chunks = []
        
        for chunk in chunks:
            content_hash = hashlib.md5(chunk.page_content.strip().encode()).hexdigest()
            if content_hash not in seen_hashes and len(chunk.page_content.strip()) > 20:
                seen_hashes.add(content_hash)
                unique_chunks.append(chunk)
        
        return unique_chunks
